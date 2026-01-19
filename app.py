import streamlit as st
import google.generativeai as genai
import json
import io
from docx import Document

# Intentamos importar el prompt, si falla usamos uno por defecto para evitar errores
try:
    from organon_prompts import SYSTEM_PROMPT_ZH as SYSTEM_PROMPT 
except ImportError:
    SYSTEM_PROMPT = "You are a legal assistant."

# --- 1. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="Organon Iudiciale | 法律書狀生成系統",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 2. ESTILO "GRAVITAS" (CSS) ---
st.markdown("""
    <style>
    /* Fondo General */
    .stApp {
        background-color: #F8F9FA;
    }
    /* Títulos Principales - Azul Marino Institucional */
    h1, h2, h3 {
        color: #1C2E4A !important;
        font-family: 'Microsoft JhengHei', '微軟正黑體', sans-serif;
        font-weight: 600;
    }
    /* Botones */
    .stButton>button {
        color: #1C2E4A;
        border: 1px solid #1C2E4A;
        border-radius: 4px;
        background-color: white;
    }
    /* Botón Primario (Generar) */
    div.stButton > button:first-child {
        background-color: #1C2E4A;
        color: white;
        border: none;
        font-weight: bold;
        box-shadow: 0 2px 5px rgba(0,0,0,0.2);
    }
    div.stButton > button:hover {
        background-color: #2c456b;
        color: white;
    }
    /* Inputs */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea {
        background-color: #FFFFFF;
        border: 1px solid #CED4DA;
        color: #212529;
    }
    /* Sidebar */
    section[data-testid="stSidebar"] {
        background-color: #E9ECEF;
    }
    /* Warning Box Customization */
    .stAlert {
        border: 1px solid #ffc107;
        background-color: #fff3cd;
        color: #856404;
    }
    </style>
""", unsafe_allow_html=True)

# --- 3. CONFIGURACIÓN GEMINI ---
MODEL_NAME = "gemini-2.0-flash" 

try:
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
except Exception:
    st.error("⚠️ 嚴重錯誤：未檢測到 GOOGLE_API_KEY。請檢查 Secrets 設定。")
    st.stop()

# --- ESTADO DE SESIÓN ---
if 'api_calls' not in st.session_state:
    st.session_state.api_calls = 0
MAX_CALLS = 10

# --- FUNCIÓN WORD ---
def crear_documento_word(titulo, cuerpo, analisis, receptor):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'PMingLiU' # MingLiU (Estándar Taiwán)
    
    doc.add_paragraph(f"致 (To): {receptor}", style='Heading 2')
    doc.add_heading(titulo, 0)
    
    doc.add_heading('書狀內容 (Content):', level=1)
    doc.add_paragraph(cuerpo)
    
    doc.add_page_break()
    doc.add_heading('附件：AI 策略分析 (Inventio Analysis)', level=1)
    doc.add_paragraph(f"爭點狀態 (Status): {analisis.get('status_causae', 'N/A')}")
    doc.add_paragraph(f"防禦策略 (Strategy): {analisis.get('estrategia_defensa', 'N/A')}")
    doc.add_paragraph(f"核心法源 (Legal Basis): {analisis.get('puntos_clave', 'N/A')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- FUNCIÓN API ---
def generar_escrito(datos):
    if st.session_state.api_calls >= MAX_CALLS:
        return {"error": "已達到試用版次數限制 (10/10)。(Demo limit reached)"}
    
    generation_config = {
        "temperature": 0.4,
        "top_p": 0.95,
        "max_output_tokens": 8192,
        "response_mime_type": "application/json",
    }

    model = genai.GenerativeModel(
        model_name=MODEL_NAME,
        system_instruction=SYSTEM_PROMPT,
        generation_config=generation_config,
    )

    user_prompt = f"""
    請根據以下資訊撰寫法律書狀 (Draft request):

    --- 基本設定 (Settings) ---
    【致送機關 (Recipient)】: {datos['receptor']}
    【語氣風格 (Tone)】: {datos['tono']} 
    (Instruction: Strictly adapt the writing style to this tone.)

    --- 案件內容 (Case Details) ---
    1. 【案情事實 (Hechos)】: 
    {datos['hechos']}
    
    2. 【引用法條 (Leyes)】: 
    {datos['leyes'] if datos['leyes'] else "由系統自行判斷適用法條"}

    3. 【引用實務見解 (Jurisprudencia)】: 
    {datos['jurisprudencia'] if datos['jurisprudencia'] else "無特定引用"}

    4. 【關鍵證據 (Pruebas)】: 
    {datos['pruebas']}
    
    5. 【對造主張 (Contraparte)】: 
    {datos['contraparte']}
    
    6. 【訴之聲明/目標 (Objetivo)】: 
    {datos['objetivo']}
    """

    try:
        response = model.generate_content(user_prompt)
        st.session_state.api_calls += 1
        return json.loads(response.text)
    except Exception as e:
        return {"error": f"系統發生錯誤 (System Error): {str(e)}"}

# --- FRONTEND ---

st.markdown("""
    <div style='text-align: center; padding-bottom: 20px;'>
        <h1 style='color: #1C2E4A; margin-bottom: 0;'>⚖️ Organon Iudiciale</h1>
        <p style='color: #666; font-size: 1.1em;'>台灣法律書狀 AI 生成系統 (專業版)</p>
        <div style='height: 2px; background-color: #1C2E4A; width: 100px; margin: 10px auto;'></div>
    </div>
""", unsafe_allow_html=True)

# --- NUEVO DESCARGO DE RESPONSABILIDAD (DISCLAIMER) ---
st.warning("""
    **⚠️ 免責聲明 (Disclaimer)**：
    本應用程式僅提供 **AI 輔助撰寫 (AI-assisted drafting)** 功能，不代表對法律之解釋或適用建議。
    生成內容僅供參考，可能存在錯誤或過時資訊，**請務必經由專業律師 (Professional Lawyer) 審閱、修改後方可使用**。
    使用者應自行確認內容之正確性與適法性。
""")
# -----------------------------------------------------

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.markdown("### 📝 案件資訊輸入 (Input)")
    st.info("💡 提示：輸入越精確的「法條」與「判決字號」，生成的書狀將越具說服力。")
    
    with st.form("main_form"):
        # Contexto
        st.markdown("**0. 致送對象與風格 (Context)**")
        c1, c2 = st.columns(2)
        with c1:
            receptor = st.text_input("受文者 (Recipient)", placeholder="例如：臺灣臺北地方法院民事庭")
        with c2:
            tono = st.text_input("書寫語氣 (Tone)", placeholder="例如：莊重保守、犀利...")

        # Hechos
        st.markdown("**1. 基礎事實與目標 (Facts)**")
        # CORRECCIÓN AQUÍ: Aseguramos que todo esté en una sola línea
        hechos = st.text_area("案情事實", height=100, placeholder="請依時間序列描述發生經過...")
        objetivo = st.text_input("訴之聲明 / 目標", placeholder="例如：請求駁回原告之訴...")

        # Leyes
        st.markdown("**2. 法源依據 (Legal Basis)**")
        leyes = st.text_area("引用法條", height=70, placeholder="例如：民法第184條...")
        with st.expander("進階：引用實務見解 (Jurisprudence)"):
            jurisprudencia = st.text_area("相關判決字號", height=70)

        # Estrategia
        st.markdown("**3. 攻防細節 (Details)**")
        pruebas = st.text_area("關鍵證據", height=70)
        contraparte = st.text_area("對造主張", height=70)
        
        submitted = st.form_submit_button("🚀 生成法律書狀 (Generate)")

with col2:
    st.markdown("### 📄 書狀預覽 (Preview)")
    
    if submitted:
        if not hechos or not objetivo or not receptor:
            st.warning("⚠️ 請填寫必要欄位：【受文者】、【案情事實】與【訴之聲明】。")
        else:
            with st.spinner(f"⚖️ Organon 正在思考中... (Engine: {MODEL_NAME})"):
                tono_final = tono if tono else "專業、莊重"
                
                datos = {
                    "receptor": receptor, "tono": tono_final,
                    "hechos": hechos, "leyes": leyes, "jurisprudencia": jurisprudencia,
                    "pruebas": pruebas, "contraparte": contraparte, "objetivo": objetivo
                }
                resultado = generar_escrito(datos)
            
            if "error" in resultado:
                st.error(f"❌ {resultado['error']}")
            else:
                doc_final = resultado.get("documento_final", {})
                analisis = resultado.get("analisis_estrategico", {})
                
                # Visualización
                with st.expander("🧠 AI 策略分析 (Strategy)", expanded=True):
                    st.markdown(f"**核心爭點:** {analisis.get('status_causae')}")
                    st.markdown(f"**防禦策略:** {analisis.get('estrategia_defensa')}")

                titulo = doc_final.get('titulo', '法律書狀')
                texto = doc_final.get('texto_completo', '')
                
                st.markdown(f"#### {titulo}")
                st.markdown(f"**致：{receptor}**") 
                st.code(texto, language=None)
                
                # Descarga
                st.markdown("---")
                docx = crear_documento_word(titulo, texto, analisis, receptor)
                st.download_button(
                    label="💾 下載 Word 檔 (.docx)",
                    data=docx,
                    file_name=f"{titulo}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

st.divider()
st.caption(f"系統狀態：Online | Model: {MODEL_NAME} (Stable) | Calls: {st.session_state.api_calls}/{MAX_CALLS}")
